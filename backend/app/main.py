import os
import tempfile
import json
import requests
from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from typing import List, Optional
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

from .models import UploadResponse, ResumeData
from .services.resume_parser import ResumeParser
from .services.data_storage import DataStorage

# 创建FastAPI应用
app = FastAPI(
    title="简历上传解析系统",
    description="支持Word格式简历的上传、解析和数据存储",
    version="1.0.0"
)

# 配置CORS - 允许所有来源（开发环境）
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 开发环境允许所有来源
    allow_credentials=False,  # allow_origins=["*"]时必须设为False
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# 表单数据模型
class ResumeFormData(BaseModel):
    姓名: Optional[str] = ""
    性别: Optional[str] = ""
    年龄: Optional[str] = ""
    政治面貌: Optional[str] = ""
    体重: Optional[str] = ""
    籍贯: Optional[str] = ""
    健康状况: Optional[str] = ""
    身高: Optional[str] = ""
    学历: Optional[str] = ""
    毕业院校: Optional[str] = ""
    专业: Optional[str] = ""
    求职意向: Optional[str] = ""
    手机: Optional[str] = ""
    邮箱: Optional[str] = ""
    教育经历: Optional[str] = ""
    荣誉奖项: Optional[str] = ""
    技能证书: Optional[str] = ""
    工作经历: Optional[str] = ""
    兴趣爱好: Optional[str] = ""
    自我评价: Optional[str] = ""

# 初始化服务
resume_parser = ResumeParser()
# 使用绝对路径确保在任何工作目录下都能找到CSV文件
import os
# 获取项目根目录（backend的上级目录）
project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
csv_path = os.path.join(project_root, "data", "resume.csv")
data_storage = DataStorage(csv_path)


@app.get("/")
async def root():
    """根路径"""
    return {"message": "简历上传解析系统 API", "version": "1.0.0"}


@app.post("/upload", response_model=UploadResponse)
async def upload_resume(file: UploadFile = File(...)):
    """
    上传并解析简历文件
    
    支持的文件格式：
    - .docx (Word文档)
    - .pdf (PDF文档)
    """
    
    # 验证文件类型
    allowed_extensions = ['.docx', '.pdf']
    file_extension = None
    for ext in allowed_extensions:
        if file.filename.lower().endswith(ext):
            file_extension = ext
            break
    
    if not file_extension:
        raise HTTPException(
            status_code=400, 
            detail="仅支持 .docx 和 .pdf 格式的文件"
        )
    
    try:
        # 创建临时文件保存上传的文件，使用正确的扩展名
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # 解析简历内容
        resume_data = resume_parser.parse_resume(temp_file_path)
        
        # 保存解析结果到CSV
        save_success = data_storage.save_resume_data(resume_data)
        
        # 清理临时文件
        os.unlink(temp_file_path)
        
        if save_success:
            return UploadResponse(
                success=True,
                message="简历上传和解析成功",
                data=resume_data
            )
        else:
            return UploadResponse(
                success=False,
                message="简历解析成功，但保存到数据库失败",
                data=resume_data
            )
            
    except Exception as e:
        # 清理临时文件
        if 'temp_file_path' in locals():
            try:
                os.unlink(temp_file_path)
            except:
                pass
        
        raise HTTPException(
            status_code=500,
            detail=f"处理文件时出错: {str(e)}"
        )


@app.get("/resumes")
async def get_resumes(
    keyword: Optional[str] = Query(None, description="搜索关键词"),
    min_age: Optional[int] = Query(None, description="最小年龄"),
    max_age: Optional[int] = Query(None, description="最大年龄"),
    min_work_years: Optional[int] = Query(None, description="最小工作年限"),
    max_work_years: Optional[int] = Query(None, description="最大工作年限"),
    limit: int = Query(100, description="返回结果数量限制")
):
    """
    获取简历列表
    
    参数：
    - keyword: 可选的搜索关键词
    - min_age: 最小年龄筛选
    - max_age: 最大年龄筛选
    - min_work_years: 最小工作年限筛选
    - max_work_years: 最大工作年限筛选
    - limit: 返回结果数量限制
    """
    try:
        # 使用新的筛选方法
        resumes = data_storage.filter_resumes(
            keyword=keyword or '',
            min_age=min_age,
            max_age=max_age,
            min_work_years=min_work_years,
            max_work_years=max_work_years
        )
        
        # 限制返回数量
        if limit > 0:
            resumes = resumes[:limit]
        
        return {
            "success": True,
            "count": len(resumes),
            "data": resumes
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"获取简历数据时出错: {str(e)}"
        )


@app.get("/stats")
async def get_stats():
    """获取系统统计信息"""
    try:
        # 获取所有简历数据进行分析
        all_resumes = data_storage.get_all_resumes()
        total_count = len(all_resumes)
        
        # 基础统计
        stats = {
            "total_resumes": total_count,
            "system_status": "运行正常"
        }
        
        if total_count > 0:
            # 性别分布统计
            gender_stats = {}
            for resume in all_resumes:
                gender = resume.get('性别', '未知')
                gender_stats[gender] = gender_stats.get(gender, 0) + 1
            stats["gender_distribution"] = gender_stats
            
            # 学历分布统计
            education_stats = {}
            for resume in all_resumes:
                education = resume.get('学历', '未知')
                education_stats[education] = education_stats.get(education, 0) + 1
            stats["education_distribution"] = education_stats
            
            # 年龄分布统计
            age_groups = {"20-25": 0, "26-30": 0, "31-35": 0, "36-40": 0, "41-45": 0, "46+": 0}
            for resume in all_resumes:
                try:
                    age = int(resume.get('年龄', 0))
                    if 20 <= age <= 25:
                        age_groups["20-25"] += 1
                    elif 26 <= age <= 30:
                        age_groups["26-30"] += 1
                    elif 31 <= age <= 35:
                        age_groups["31-35"] += 1
                    elif 36 <= age <= 40:
                        age_groups["36-40"] += 1
                    elif 41 <= age <= 45:
                        age_groups["41-45"] += 1
                    elif age >= 46:
                        age_groups["46+"] += 1
                except (ValueError, TypeError):
                    pass
            stats["age_distribution"] = age_groups
            
            # 热门专业TOP10
            major_stats = {}
            for resume in all_resumes:
                major = resume.get('专业', '')
                if major and major != '':
                    major_stats[major] = major_stats.get(major, 0) + 1
            
            # 按数量排序，取前10
            sorted_majors = sorted(major_stats.items(), key=lambda x: x[1], reverse=True)[:10]
            stats["top_majors"] = dict(sorted_majors)
            
            # 热门院校TOP10
            school_stats = {}
            for resume in all_resumes:
                school = resume.get('毕业院校', '')
                if school and school != '':
                    school_stats[school] = school_stats.get(school, 0) + 1
            
            sorted_schools = sorted(school_stats.items(), key=lambda x: x[1], reverse=True)[:10]
            stats["top_schools"] = dict(sorted_schools)
            
            # 热门求职意向TOP10
            job_intention_stats = {}
            for resume in all_resumes:
                job_intention = resume.get('求职意向', '')
                if job_intention and job_intention != '':
                    job_intention_stats[job_intention] = job_intention_stats.get(job_intention, 0) + 1
            
            sorted_job_intentions = sorted(job_intention_stats.items(), key=lambda x: x[1], reverse=True)[:10]
            stats["top_job_intentions"] = dict(sorted_job_intentions)
            
            # 工作年限分布
            work_years_groups = {"0年": 0, "1-3年": 0, "4-6年": 0, "7-10年": 0, "11-15年": 0, "15年+": 0}
            for resume in all_resumes:
                work_years = data_storage._calculate_work_years(
                    resume.get('工作经历', ''), 
                    resume.get('年龄', '')
                )
                if work_years == 0:
                    work_years_groups["0年"] += 1
                elif 1 <= work_years <= 3:
                    work_years_groups["1-3年"] += 1
                elif 4 <= work_years <= 6:
                    work_years_groups["4-6年"] += 1
                elif 7 <= work_years <= 10:
                    work_years_groups["7-10年"] += 1
                elif 11 <= work_years <= 15:
                    work_years_groups["11-15年"] += 1
                else:
                    work_years_groups["15年+"] += 1
            stats["work_years_distribution"] = work_years_groups
            
            # 录入时间分析
            from datetime import datetime, timedelta
            today = datetime.now().date()
            time_groups = {"今天": 0, "本周": 0, "本月": 0, "更早": 0}
            
            for resume in all_resumes:
                try:
                    record_time_str = resume.get('录入时间', '')
                    if record_time_str:
                        record_time = datetime.strptime(record_time_str, '%Y-%m-%d %H:%M:%S').date()
                        
                        if record_time == today:
                            time_groups["今天"] += 1
                        elif record_time >= today - timedelta(days=7):
                            time_groups["本周"] += 1
                        elif record_time >= today.replace(day=1):
                            time_groups["本月"] += 1
                        else:
                            time_groups["更早"] += 1
                except:
                    time_groups["更早"] += 1
            
            stats["time_distribution"] = time_groups
        
        return {
            "success": True,
            "data": stats
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"获取统计信息时出错: {str(e)}"
        )


@app.post("/submit-form")
async def submit_resume_form(form_data: ResumeFormData):
    """
    提交简历表单数据
    
    接收前端表单提交的简历数据，直接保存到CSV文件
    """
    try:
        # 将表单数据转换为ResumeData格式
        resume_data = ResumeData(
            姓名=form_data.姓名 or "",
            性别=form_data.性别 or "",
            年龄=form_data.年龄 or "",
            政治面貌=form_data.政治面貌 or "",
            体重=form_data.体重 or "",
            籍贯=form_data.籍贯 or "",
            健康状况=form_data.健康状况 or "",
            身高=form_data.身高 or "",
            学历=form_data.学历 or "",
            毕业院校=form_data.毕业院校 or "",
            专业=form_data.专业 or "",
            求职意向=form_data.求职意向 or "",
            手机=form_data.手机 or "",
            邮箱=form_data.邮箱 or "",
            教育经历=form_data.教育经历 or "",
            荣誉奖项=form_data.荣誉奖项 or "",
            技能证书=form_data.技能证书 or "",
            工作经历=form_data.工作经历 or "",
            兴趣爱好=form_data.兴趣爱好 or "",
            自我评价=form_data.自我评价 or ""
        )
        
        # 保存到CSV文件
        save_success = data_storage.save_resume_data(resume_data)
        
        if save_success:
            return {
                "success": True,
                "message": "简历表单提交成功",
                "data": resume_data.dict()
            }
        else:
            raise HTTPException(
                status_code=500,
                detail="保存简历数据失败"
            )
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"提交简历表单时出错: {str(e)}"
        )


@app.post("/generate-resume-data")
async def generate_resume_data():
    """
    使用DeepSeek AI生成简历测试数据
    """
    try:
        # DeepSeek API配置
        api_key = "sk-9b31446d564c46a2b1593be7804f4376"
        api_url = "https://api.deepseek.com/v1/chat/completions"
        
        # 添加时间戳和随机元素来增加随机性
        import time
        import random
        
        timestamp = int(time.time())
        random_num = random.randint(1000, 9999)
        
        # 随机选择不同的专业领域和背景
        majors = ["计算机科学与技术", "软件工程", "信息安全", "数据科学与大数据技术", "人工智能", "电子信息工程", "通信工程", "自动化", "机械工程", "电气工程", "市场营销", "会计学", "人力资源管理", "国际贸易", "英语", "汉语言文学", "新闻学", "广告学", "工商管理", "金融学"]
        
        # 混合知名公司和普通公司
        companies = [
            # 知名互联网公司
            "阿里巴巴", "腾讯", "百度", "字节跳动", "美团", "滴滴", "京东", "网易", "小米", "华为",
            # 普通互联网公司
            "北京创新科技有限公司", "上海智慧信息技术有限公司", "深圳天翼软件开发有限公司", "杭州云端科技有限公司",
            # 传统行业公司
            "中国建设银行", "工商银行", "招商银行", "中国移动", "中国电信", "国家电网", "中石化", "中石油",
            # 中小企业
            "蓝天贸易有限公司", "星辰教育科技有限公司", "绿叶环保科技有限公司", "金桥广告有限公司", "新华书店", "华联超市",
            # 外企
            "微软中国", "IBM中国", "英特尔中国", "西门子中国", "宝马中国", "奔驰中国"
        ]
        
        cities = ["北京", "上海", "深圳", "杭州", "广州", "成都", "西安", "南京", "武汉", "苏州", "天津", "重庆", "青岛", "大连", "厦门", "宁波", "无锡", "长沙", "郑州", "济南"]
        
        # 混合985/211和普通院校
        universities = [
            # 985/211院校
            "清华大学", "北京大学", "复旦大学", "上海交通大学", "浙江大学", "华中科技大学", "西安交通大学", "哈尔滨工业大学",
            "中山大学", "华南理工大学", "北京理工大学", "东南大学", "电子科技大学", "西北工业大学",
            # 普通一本院校  
            "北京工业大学", "上海理工大学", "广东工业大学", "河北工业大学", "南京理工大学", "西安理工大学",
            "大连理工大学", "长春理工大学", "成都理工大学", "桂林理工大学", "昆明理工大学", "太原理工大学",
            # 普通二本院校
            "河北科技大学", "山东建筑大学", "湖南工程学院", "江西理工大学", "安徽工程大学", "河南工程学院",
            "湖北工程学院", "广西科技大学", "云南农业大学", "贵州财经大学", "新疆师范大学", "内蒙古工业大学",
            # 职业院校
            "北京信息职业技术学院", "上海电子信息职业技术学院", "深圳职业技术学院", "广州番禺职业技术学院"
        ]
        
        # 增加随机姓名池，确保姓名多样性
        surnames = ["王", "李", "张", "刘", "陈", "杨", "黄", "赵", "周", "吴", "徐", "孙", "马", "朱", "胡", "林", "郭", "何", "高", "罗"]
        male_names = ["明", "强", "军", "杰", "涛", "磊", "勇", "伟", "峰", "飞", "辉", "建", "华", "鹏", "宇", "浩", "凯", "俊", "超", "晨"]
        female_names = ["丽", "娜", "敏", "静", "秀", "雪", "梅", "红", "玲", "芳", "晴", "欣", "瑶", "莹", "萍", "燕", "霞", "倩", "婷", "洁"]
        
        selected_major = random.choice(majors)
        selected_company = random.choice(companies)
        selected_city = random.choice(cities)
        selected_university = random.choice(universities)
        selected_surname = random.choice(surnames)
        
        # 随机选择性别和对应的名字
        gender = random.choice(["男", "女"])
        if gender == "男":
            selected_name = selected_surname + random.choice(male_names) + (random.choice(male_names) if random.random() > 0.5 else "")
        else:
            selected_name = selected_surname + random.choice(female_names) + (random.choice(female_names) if random.random() > 0.5 else "")
        
        # 生成随机手机号
        phone_prefixes = ["138", "139", "150", "151", "152", "158", "159", "130", "131", "132", "155", "156", "186", "185", "176"]
        selected_phone = random.choice(phone_prefixes) + str(random.randint(10000000, 99999999))
        
        # 随机生成年龄（20-50岁）
        age = random.randint(20, 50)
        
        # 根据年龄调整工作经历段数和逻辑
        if age <= 23:
            # 20-23岁：应届生或刚工作1-2年
            work_experience_count = random.choice([1, 1, 2])  # 至少1段经历（可能是实习）
            work_description = "应届生或刚参加工作，可能主要是实习经历"
        elif age <= 27:
            # 24-27岁：工作2-5年
            work_experience_count = random.choice([1, 2, 2])
            work_description = "工作2-5年，有一定工作经验"
        elif age <= 32:
            # 28-32岁：工作5-10年
            work_experience_count = random.choice([2, 3, 3])
            work_description = "工作5-10年，有丰富工作经验"
        elif age <= 40:
            # 33-40岁：工作10-18年
            work_experience_count = random.choice([3, 4, 4])
            work_description = "工作10-18年，资深从业者"
        else:
            # 41-50岁：工作18年以上
            work_experience_count = random.choice([4, 5, 5])
            work_description = "工作18年以上，资深专家或管理者"
        
        # 构建简化的AI提示词，提升生成速度
        graduation_year = 2024 - (age - 22)
        education_start_year = graduation_year - 4
        work_start_year = graduation_year
        
        # 生成随机邮箱
        email_domains = ["qq.com", "163.com", "126.com", "gmail.com", "sina.com", "hotmail.com"]
        selected_email = f"{selected_name.lower()}{random.randint(100, 999)}@{random.choice(email_domains)}"
        
        prompt = f"""生成中文简历JSON数据。时间戳：{timestamp}

指定参数：
- 姓名：{selected_name}
- 性别：{gender}
- 年龄：{age}岁
- 手机：{selected_phone}
- 专业：{selected_major}
- 院校：{selected_university}
- 工作公司：{selected_company}

要求：
1. 生成{work_experience_count}段工作经历
2. {work_description}
3. 本科毕业时间约{graduation_year}年
4. 工作经历从毕业后开始，时间连贯
5. 包含{selected_company}公司
6. 直接返回JSON格式，无其他文字

JSON结构示例：
{{
  "姓名": "{selected_name}",
  "性别": "{gender}",
  "年龄": {age},
  "政治面貌": "共青团员",
  "体重": "65kg",
  "籍贯": "江苏省南京市",
  "健康状况": "良好",
  "身高": "170cm",
  "学历": "本科",
  "毕业院校": "{selected_university}",
  "专业": "{selected_major}",
  "求职意向": "相关职位名称",
  "手机": "{selected_phone}",
  "邮箱": "{selected_email}",
  "教育经历": [
    {{
      "start_time": "{education_start_year}-09",
      "end_time": "{graduation_year}-06",
      "school": "{selected_university}",
      "degree": "本科",
      "major": "{selected_major}"
    }}
  ],
  "荣誉奖项": "相关专业奖项",
  "技能证书": "相关专业证书",
  "工作经历": [
    {{
      "start_time": "{work_start_year}-07",
      "end_time": "至今",
      "company": "{selected_company}",
      "position": "相关职位",
      "responsibilities": "工作职责描述"
    }}
  ],
  "兴趣爱好": "运动、阅读、音乐",
  "自我评价": "简短专业评价"
}}

请严格按照以上JSON结构生成数据：
1. 教育经历必须是数组格式，包含完整的教育信息
2. 工作经历必须是数组格式，包含{work_experience_count}段不同的工作记录
3. 工作经历时间要连贯，从{work_start_year}年开始，职位要有发展路径
4. 其中一段工作经历必须在{selected_company}公司
5. 所有时间格式为"YYYY-MM"，最新工作可用"至今"
6. 直接返回JSON，不要添加任何解释文字"""

        # 调用DeepSeek API
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.9,  # 提高温度以增加随机性
            "top_p": 0.95,       # 添加top_p参数增加多样性
            "max_tokens": 1500   # 减少token数量，提升生成速度
        }
        
        # 添加重试机制，提升稳定性
        max_retries = 3
        last_error = None
        
        for attempt in range(max_retries):
            try:
                print(f"AI生成尝试 {attempt + 1}/{max_retries}")
                
                response = requests.post(
                    api_url, 
                    headers=headers, 
                    json=data,
                    timeout=50  # 设置50秒超时
                )
                
                if response.status_code != 200:
                    if attempt < max_retries - 1:
                        print(f"API调用失败，状态码: {response.status_code}，重试中...")
                        last_error = f"API status code: {response.status_code}"
                        continue
                    raise HTTPException(
                        status_code=500,
                        detail=f"DeepSeek API调用失败: {response.status_code}"
                    )
                
                result = response.json()
                
                if 'choices' not in result or not result['choices']:
                    if attempt < max_retries - 1:
                        print("API响应格式异常，重试中...")
                        last_error = "Invalid API response format"
                        continue
                    raise HTTPException(
                        status_code=500,
                        detail="API响应格式异常"
                    )
                
                ai_content = result["choices"][0]["message"]["content"]
                print(f"AI生成成功，内容长度: {len(ai_content)}")
                print(f"AI原始内容: {ai_content[:500]}...")  # 打印前500字符用于调试
                break  # 成功，跳出重试循环
                
            except requests.exceptions.Timeout:
                if attempt < max_retries - 1:
                    print(f"请求超时，重试中... ({attempt + 1}/{max_retries})")
                    last_error = "Request timeout"
                    continue
                raise HTTPException(
                    status_code=408,
                    detail="AI生成超时，请稍后重试"
                )
            except requests.exceptions.RequestException as e:
                if attempt < max_retries - 1:
                    print(f"网络请求失败，重试中... 错误: {e}")
                    last_error = str(e)
                    continue
                raise HTTPException(
                    status_code=502,
                    detail=f"网络请求失败: {str(e)}"
                )
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"其他错误，重试中... 错误: {e}")
                    last_error = str(e)
                    continue
                raise HTTPException(
                    status_code=500,
                    detail=f"AI调用失败: {str(e)}"
                )
        else:
            # 如果所有重试都失败了
            raise HTTPException(
                status_code=500,
                detail=f"AI生成失败，已重试{max_retries}次。最后错误: {last_error}"
            )
        
        # 尝试解析AI返回的JSON数据
        try:
            # 提取JSON部分（如果AI返回了额外的文本）
            import re
            json_match = re.search(r'\{.*\}', ai_content, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                generated_data = json.loads(json_str)
            else:
                # 如果没有找到JSON，尝试直接解析
                generated_data = json.loads(ai_content)
            
            # 验证生成的数据结构
            print(f"验证生成数据 - 姓名: {generated_data.get('姓名', 'N/A')}")
            print(f"教育经历类型: {type(generated_data.get('教育经历', None))}, 内容: {generated_data.get('教育经历', [])}")
            print(f"工作经历类型: {type(generated_data.get('工作经历', None))}, 长度: {len(generated_data.get('工作经历', []))}")
            
            # 确保教育经历和工作经历是数组格式
            if not isinstance(generated_data.get('教育经历'), list):
                print("警告：教育经历不是数组格式，使用默认数据")
                generated_data['教育经历'] = [{
                    "start_time": f"{education_start_year}-09",
                    "end_time": f"{graduation_year}-06",
                    "school": selected_university,
                    "degree": "本科",
                    "major": selected_major
                }]
            
            if not isinstance(generated_data.get('工作经历'), list) or len(generated_data.get('工作经历', [])) == 0:
                print("警告：工作经历不是数组格式或为空，使用默认数据")
                generated_data['工作经历'] = [{
                    "start_time": f"{work_start_year}-07",
                    "end_time": "至今",
                    "company": selected_company,
                    "position": "相关职位",
                    "responsibilities": "负责相关工作内容"
                }]
                
        except json.JSONDecodeError:
            # 如果解析失败，返回默认测试数据
            generated_data = {
                "姓名": "张小明",
                "性别": "男",
                "年龄": "26",
                "政治面貌": "共青团员",
                "体重": "68",
                "籍贯": "北京市",
                "健康状况": "良好",
                "身高": "175",
                "学历": "硕士",
                "毕业院校": "清华大学",
                "专业": "计算机科学与技术",
                "求职意向": "软件开发工程师",
                "手机": "138-0013-8000",
                "邮箱": "zhangxiaoming@example.com",
                "教育经历": [
                    {
                        "start_time": "2018-09",
                        "end_time": "2020-06",
                        "school": "清华大学",
                        "degree": "硕士",
                        "major": "计算机科学与技术"
                    },
                    {
                        "start_time": "2014-09",
                        "end_time": "2018-06",
                        "school": "北京大学",
                        "degree": "本科",
                        "major": "软件工程"
                    }
                ],
                "荣誉奖项": "国家奖学金、ACM程序设计竞赛一等奖、优秀毕业生",
                "技能证书": "Java开发工程师认证、AWS云计算认证、PMP项目管理认证",
                "工作经历": [
                    {
                        "start_time": "2020-07",
                        "end_time": "至今",
                        "company": "阿里巴巴集团",
                        "position": "高级软件开发工程师",
                        "responsibilities": "负责电商平台后端服务开发，参与微服务架构设计，优化系统性能提升30%，带领5人技术团队完成多个核心项目"
                    },
                    {
                        "start_time": "2020-03",
                        "end_time": "2020-06",
                        "company": "腾讯科技（实习）",
                        "position": "软件开发实习生",
                        "responsibilities": "参与微信小程序开发，协助完成用户画像系统，学习大数据处理技术"
                    }
                ],
                "兴趣爱好": "编程、阅读、篮球、摄影、旅行",
                "自我评价": "具有扎实的计算机专业基础和丰富的项目开发经验，熟练掌握Java、Python等编程语言，对分布式系统和微服务架构有深入理解。具备良好的团队合作能力和技术领导力，善于学习新技术并能快速应用到实际项目中。对软件工程有浓厚兴趣，致力于编写高质量、可维护的代码。"
            }
        
        return {
            "success": True,
            "message": "AI生成简历数据成功",
            "data": generated_data
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"生成简历数据时出错: {str(e)}"
        )


@app.delete("/resumes/{unique_id}")
async def delete_resume(unique_id: str):
    """
    删除指定唯一标识符的简历记录
    
    参数：
    - unique_id: 简历唯一标识符（姓名_录入时间）
    """
    try:
        # URL解码
        import urllib.parse
        decoded_id = urllib.parse.unquote(unique_id)
        
        success = data_storage.delete_resume_by_id(decoded_id)
        
        if success:
            return {
                "success": True,
                "message": "简历删除成功"
            }
        else:
            raise HTTPException(
                status_code=404,
                detail="简历记录不存在或删除失败"
            )
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"删除简历时出错: {str(e)}"
        )


# Word导出请求数据模型
class WordExportRequest(BaseModel):
    resume: dict

@app.post("/export-word")
async def export_word(request: WordExportRequest):
    """
    导出简历为Word文档
    
    接收简历数据，生成Word格式的简历文档并返回文件
    """
    try:
        resume_data = request.resume
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading('个人简历', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 安全地获取和转换所有字段为字符串
        def safe_str(value):
            if value is None:
                return ''
            return str(value)
        
        # 添加空行
        doc.add_paragraph('')
        
        # 创建基本信息表格 - 使用简单的2列布局（标签 + 内容）
        basic_info_table = doc.add_table(rows=12, cols=2)
        basic_info_table.style = 'Table Grid'
        
        # 基本信息数据 - 每行一个信息项
        basic_info_items = [
            ('姓名', safe_str(resume_data.get('姓名', ''))),
            ('性别', safe_str(resume_data.get('性别', ''))),
            ('年龄', safe_str(resume_data.get('年龄', ''))),
            ('政治面貌', safe_str(resume_data.get('政治面貌', ''))),
            ('体重', safe_str(resume_data.get('体重', ''))),
            ('籍贯', safe_str(resume_data.get('籍贯', ''))),
            ('健康状况', safe_str(resume_data.get('健康状况', ''))),
            ('身高', safe_str(resume_data.get('身高', ''))),
            ('学历', safe_str(resume_data.get('学历', ''))),
            ('毕业院校', safe_str(resume_data.get('毕业院校', ''))),
            ('专业', safe_str(resume_data.get('专业', ''))),
            ('求职意向', safe_str(resume_data.get('求职意向', '')))
        ]
        
        # 填充基本信息表格
        for row_idx, (label, value) in enumerate(basic_info_items):
            if row_idx < len(basic_info_table.rows):
                row = basic_info_table.rows[row_idx]
                # 标签列
                label_cell = row.cells[0]
                label_cell.text = label
                # 设置标签为粗体
                for paragraph in label_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                
                # 内容列
                content_cell = row.cells[1]
                content_cell.text = value
        
        # 添加联系信息表格
        doc.add_paragraph('')
        contact_table = doc.add_table(rows=2, cols=2)
        contact_table.style = 'Table Grid'
        
        # 联系信息
        contact_items = [
            ('手机', safe_str(resume_data.get('手机', ''))),
            ('邮箱', safe_str(resume_data.get('邮箱', '')))
        ]
        
        for row_idx, (label, value) in enumerate(contact_items):
            if row_idx < len(contact_table.rows):
                row = contact_table.rows[row_idx]
                # 标签列
                label_cell = row.cells[0]
                label_cell.text = label
                # 设置标签为粗体
                for paragraph in label_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                
                # 内容列
                content_cell = row.cells[1]
                content_cell.text = value
        
        # 添加详细信息部分
        sections = [
            ('教育经历', resume_data.get('教育经历', '')),
            ('工作经历', resume_data.get('工作经历', '')),
            ('荣誉奖项', resume_data.get('荣誉奖项', '')),
            ('技能证书', resume_data.get('技能证书', '')),
            ('兴趣爱好', resume_data.get('兴趣爱好', '')),
            ('自我评价', resume_data.get('自我评价', ''))
        ]
        
        for section_title, section_content in sections:
            if section_content:
                # 添加分节标题
                doc.add_paragraph('')
                section_heading = doc.add_heading(section_title, level=2)
                section_heading.style.font.color.rgb = RGBColor(0, 0, 0)
                
                # 处理教育经历和工作经历（可能是JSON格式）
                if section_title in ['教育经历', '工作经历']:
                    content_text = format_json_content(section_content, section_title)
                else:
                    content_text = str(section_content)
                
                # 添加分节内容
                section_para = doc.add_paragraph(content_text)
                section_para.style = 'Normal'
        
        # 保存文档到内存
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # 生成文件名，使用URL编码处理中文
        import urllib.parse
        base_filename = f"{resume_data.get('姓名', '简历')}_简历.docx"
        encoded_filename = urllib.parse.quote(base_filename.encode('utf-8'))
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(doc_buffer.getvalue())
        temp_file.close()
        
        # 返回文件
        return FileResponse(
            temp_file.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=base_filename,
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
        
    except Exception as e:
        print(f"Word导出错误: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Word文档生成失败: {str(e)}"
        )

def format_json_content(content, content_type):
    """
    格式化JSON内容为可读文本
    """
    try:
        if isinstance(content, str):
            # 尝试解析JSON字符串
            import json
            data = json.loads(content.replace('""', '"'))
        elif isinstance(content, list):
            data = content
        else:
            return str(content)
        
        if content_type == '教育经历':
            formatted_parts = []
            for edu in data:
                parts = []
                if edu.get('start_time') and edu.get('end_time'):
                    parts.append(f"起止时间：{edu['start_time']} – {edu['end_time']}")
                if edu.get('school'):
                    parts.append(f"学校：{edu['school']}")
                if edu.get('degree'):
                    parts.append(f"学历：{edu['degree']}")
                if edu.get('major'):
                    parts.append(f"专业：{edu['major']}")
                formatted_parts.append('\n'.join(parts))
            return '\n\n'.join(formatted_parts)
            
        elif content_type == '工作经历':
            formatted_parts = []
            for work in data:
                parts = []
                if work.get('start_time') and work.get('end_time'):
                    parts.append(f"起止时间：{work['start_time']} – {work['end_time']}")
                if work.get('company'):
                    parts.append(f"公司：{work['company']}")
                if work.get('position'):
                    parts.append(f"职位：{work['position']}")
                if work.get('responsibilities'):
                    parts.append(f"主要职责：{work['responsibilities']}")
                formatted_parts.append('\n'.join(parts))
            return '\n\n'.join(formatted_parts)
        
        return str(content)
        
    except Exception as e:
        print(f"JSON解析错误: {e}")
        return str(content)

@app.get("/health")
async def health_check():
    """健康检查接口"""
    return {
        "status": "healthy",
        "message": "服务运行正常"
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080) 