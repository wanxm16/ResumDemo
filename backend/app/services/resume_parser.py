import re
import docx
import json
import pdfplumber
import easyocr
import io
import ssl
from PIL import Image
from docx import Document
from typing import Dict, Optional, List
from ..models import ResumeData

# 修复PIL.Image兼容性问题
if not hasattr(Image, 'ANTIALIAS'):
    Image.ANTIALIAS = Image.LANCZOS


class ResumeParser:
    """简历解析器"""
    
    def __init__(self):
        # 延迟初始化OCR读取器，避免启动时的SSL问题
        self.ocr_reader = None
        
        # 定义各个字段的关键词匹配规则
        self.field_patterns = {
            '姓名': [r'姓\s*名[：:]\s*([^\s\n]+)', r'姓名\s*([^\s\n]+)', r'^([^\s\n]{2,4})\s*$'],
            '性别': [r'性\s*别[：:]\s*([男女])', r'性别\s*([男女])'],
            '年龄': [r'年\s*龄[：:]\s*(\d+)', r'年龄\s*(\d+)', r'(\d+)\s*岁'],
            '政治面貌': [r'政治面貌[：:]\s*([^\s\n]+)', r'政治面貌\s*([^\s\n]+)'],
            '体重': [r'体\s*重[：:]\s*(\d+(?:\.\d+)?)\s*[公斤kgKG]?', r'(\d+(?:\.\d+)?)\s*[公斤kgKG]'],
            '籍贯': [r'籍\s*贯[：:]\s*([^\s\n]+)', r'籍贯\s*([^\s\n]+)', r'出生地[：:]\s*([^\s\n]+)'],
            '健康状况': [r'健康状况[：:]\s*([^\s\n]+)', r'健康\s*([^\s\n]+)'],
            '身高': [r'身\s*高[：:]\s*(\d+(?:\.\d+)?)\s*[厘米cmCM]?', r'(\d+(?:\.\d+)?)\s*[厘米cmCM]'],
            '学历': [r'学\s*历[：:]\s*([^\s\n]+)', r'学历\s*([^\s\n]+)', r'(本科|硕士|博士|专科|高中|初中)'],
            '毕业院校': [r'毕业院校[：:]\s*([^\s\n\|]+)', r'院校[：:]\s*([^\s\n\|]+)', r'毕业院校\s+([^\s\n\|]+)', r'([^\s\n\|]*(?:大学|学院|职业学院|技术学院)[^\s\n\|]*)'],
            '专业': [r'专\s*业[：:]\s*([^\s\n]+)', r'专业\s*([^\s\n]+)', r'专业方向[：:]\s*([^\s\n]+)'],
            '求职意向': [r'求职意向[：:]\s*([^\s\n\|]+)', r'意向职位[：:]\s*([^\s\n\|]+)', r'应聘职位[：:]\s*([^\s\n\|]+)', r'求职意向\s+([^\s\n\|]+)'],
            '手机': [r'手\s*机[：:]\s*([1][3-9]\d{1}[-\s]?\d{4}[-\s]?\d{4})', r'电话[：:]\s*([1][3-9]\d{1}[-\s]?\d{4}[-\s]?\d{4})', r'联系电话[：:]\s*([1][3-9]\d{1}[-\s]?\d{4}[-\s]?\d{4})', r'手机\s+([1][3-9]\d{1}[-\s]?\d{4}[-\s]?\d{4})', r'([1][3-9]\d{1}[-\s]?\d{4}[-\s]?\d{4})'],
            '邮箱': [
                r'(?:邮\s*箱|电子邮箱|联系邮箱|邮箱地址|电子邮件)[：:]\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'(?:Email|E-mail|email|e-mail)[：:]\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'(?:Email|E-mail|email|e-mail)[：:]([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'
            ],
        }
        
        # 定义章节关键词
        self.section_patterns = {
            '教育经历': [r'教育经历', r'教育背景', r'学习经历', r'教育情况'],
            '荣誉奖项': [r'荣誉奖项', r'获奖情况', r'奖项荣誉', r'荣誉证书'],
            '技能证书': [r'技能证书', r'专业技能', r'技能水平', r'证书情况', r'资格证书'],
            '工作经历': [r'工作经历', r'工作经验', r'职业经历', r'实习经历', r'项目经验'],
            '兴趣爱好': [r'兴趣爱好', r'个人爱好', r'业余爱好', r'特长爱好'],
            '自我评价': [r'自我评价', r'个人评价', r'自我介绍', r'个人简介', r'个人总结']
        }
    
    def parse_resume(self, file_path: str) -> ResumeData:
        """解析简历文档，支持Word和PDF格式"""
        try:
            # 根据文件扩展名判断文件类型
            if file_path.lower().endswith('.pdf'):
                text = self._extract_text_from_pdf(file_path)
            else:
                doc = Document(file_path)
                text = self._extract_text_from_doc(doc)
            
            # 解析各个字段
            parsed_data = {}
            
            # 解析基本信息字段
            for field, patterns in self.field_patterns.items():
                parsed_data[field] = self._extract_field(text, patterns)
            
            # 解析章节内容
            for section, patterns in self.section_patterns.items():
                if section == '教育经历':
                    parsed_data[section] = self._parse_education_experience(text, patterns)
                elif section == '工作经历':
                    parsed_data[section] = self._parse_work_experience(text, patterns)
                else:
                    parsed_data[section] = self._extract_section_content(text, patterns)
            
            return ResumeData(**parsed_data)
            
        except Exception as e:
            print(f"解析简历时出错: {str(e)}")
            return ResumeData()
    
    def _get_ocr_reader(self):
        """获取OCR读取器，延迟初始化"""
        if self.ocr_reader is None:
            try:
                # 处理SSL证书问题
                ssl._create_default_https_context = ssl._create_unverified_context
                print("初始化OCR读取器...")
                self.ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
                print("OCR读取器初始化成功")
            except Exception as e:
                print(f"OCR读取器初始化失败: {str(e)}")
                self.ocr_reader = None
        return self.ocr_reader
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF文件中提取文本，支持OCR识别图片内容"""
        text_parts = []
        has_images = False
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # 尝试提取文本内容
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                    
                    # 提取表格数据
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                # 过滤掉空单元格并连接
                                row_text = ' '.join([cell.strip() for cell in row if cell and cell.strip()])
                                if row_text:
                                    text_parts.append(row_text)
                    
                    # 检查是否有图片内容
                    if page.images:
                        has_images = True
                        print(f"检测到PDF第{pdf.pages.index(page)+1}页包含{len(page.images)}张图片")
                
                # 如果没有提取到文本但有图片，尝试OCR识别
                if (not text_parts or all(not part.strip() for part in text_parts)) and has_images:
                    print("开始使用OCR识别图片内容...")
                    ocr_text = self._extract_text_from_pdf_images(file_path)
                    if ocr_text:
                        text_parts.append(ocr_text)
                        print(f"OCR识别成功，提取到{len(ocr_text)}个字符")
            
            result_text = '\n'.join(text_parts)
            print(f"PDF文本提取完成，总字符数: {len(result_text)}")
            return result_text
            
        except Exception as e:
            print(f"从PDF提取文本时出错: {str(e)}")
            return ""
    
    def _extract_text_from_pdf_images(self, file_path: str) -> str:
        """使用OCR从PDF图片中提取文本"""
        try:
            # 获取OCR读取器
            ocr_reader = self._get_ocr_reader()
            if not ocr_reader:
                print("OCR读取器不可用，跳过图片识别")
                return ""
            
            ocr_text_parts = []
            
            # 尝试使用PyMuPDF提取图片（如果可用）
            try:
                import fitz  # PyMuPDF，用于更好的图片提取
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    # 将页面转换为图片
                    mat = fitz.Matrix(2.0, 2.0)  # 2倍缩放提高OCR精度
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    
                    # 使用EasyOCR识别
                    results = ocr_reader.readtext(img_data)
                    
                    page_text = []
                    for (bbox, text, confidence) in results:
                        if confidence > 0.5:  # 置信度阈值
                            page_text.append(text.strip())
                    
                    if page_text:
                        ocr_text_parts.append('\n'.join(page_text))
                        print(f"第{page_num+1}页OCR识别到{len(page_text)}个文本块")
                
                doc.close()
                
            except ImportError:
                # 如果没有PyMuPDF，使用pdfplumber的图片提取
                print("PyMuPDF不可用，使用pdfplumber进行图片提取...")
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            # 将整个页面转换为图片进行OCR
                            page_image = page.to_image(resolution=150)  # 降低分辨率避免内存问题
                            
                            # 将pdfplumber的图片转换为PIL Image并保存为字节流
                            img_bytes = io.BytesIO()
                            page_image.save(img_bytes, format='PNG')
                            img_data = img_bytes.getvalue()
                            
                            # 使用EasyOCR识别
                            results = ocr_reader.readtext(img_data)
                            
                            page_text = []
                            for (bbox, text, confidence) in results:
                                if confidence > 0.5:  # 置信度阈值
                                    page_text.append(text.strip())
                            
                            if page_text:
                                ocr_text_parts.append('\n'.join(page_text))
                                print(f"第{page_num+1}页OCR识别到{len(page_text)}个文本块")
                        except Exception as e:
                            print(f"第{page_num+1}页OCR处理失败: {str(e)}")
                            continue
            
            return '\n'.join(ocr_text_parts)
            
        except Exception as e:
            print(f"OCR识别失败: {str(e)}")
            return ""
    
    def _extract_text_from_doc(self, doc: Document) -> str:
        """从Word文档中提取文本"""
        text_parts = []
        extracted_texts = set()  # 用于去重
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and text not in extracted_texts:
                text_parts.append(text)
                extracted_texts.add(text)
        
        # 提取表格文本（只提取不重复的内容）
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in extracted_texts:
                        row_text.append(cell_text)
                        extracted_texts.add(cell_text)
                if row_text:
                    text_parts.append(' '.join(row_text))
        
        return '\n'.join(text_parts)
    
    def _extract_field(self, text: str, patterns: list) -> Optional[str]:
        """使用正则表达式提取字段"""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = match.group(1).strip()
                if value and len(value) <= 100:  # 避免提取过长的内容
                    return value
        return None
    
    def _extract_section_content(self, text: str, patterns: list) -> Optional[str]:
        """提取章节内容"""
        for pattern in patterns:
            # 查找章节标题
            section_match = re.search(f'{pattern}[：:]?', text, re.IGNORECASE)
            if section_match:
                start_pos = section_match.end()
                
                # 查找下一个章节的开始位置
                next_section_pos = len(text)
                all_section_patterns = []
                for section_list in self.section_patterns.values():
                    all_section_patterns.extend(section_list)
                
                for next_pattern in all_section_patterns:
                    if next_pattern != pattern:
                        next_match = re.search(f'{next_pattern}[：:]?', text[start_pos:], re.IGNORECASE)
                        if next_match:
                            next_section_pos = min(next_section_pos, start_pos + next_match.start())
                
                # 提取章节内容
                section_content = text[start_pos:next_section_pos].strip()
                
                # 清理内容
                section_content = re.sub(r'\n+', '\n', section_content)
                section_content = section_content[:500] if len(section_content) > 500 else section_content
                
                if section_content:
                    return section_content
        
        return None
    
    def _parse_education_experience(self, text: str, patterns: list) -> Optional[str]:
        """解析教育经历为JSON格式"""
        section_content = self._extract_section_content(text, patterns)
        if not section_content:
            # 尝试从整个文本中提取表格格式的教育经历
            return self._parse_table_education_experience(text)
        
        try:
            education_list = []
            
            # 按段落分割教育经历
            sections = re.split(r'\n\s*\n|\n(?=起止时间)', section_content)
            
            for section in sections:
                if not section.strip():
                    continue
                
                education_item = {}
                
                # 提取起止时间
                time_match = re.search(r'起止时间[：:]\s*([^\n]+)', section)
                if time_match:
                    time_range = time_match.group(1).strip()
                    # 使用更精确的正则来分割时间范围，避免分割日期中的连字符
                    # 匹配类似 "2015-09 – 2019-06" 的格式
                    time_split = re.match(r'(\d{4}-\d{2})\s*[–—-]\s*(\d{4}-\d{2}|\S+)', time_range)
                    if time_split:
                        education_item['start_time'] = time_split.group(1).strip()
                        education_item['end_time'] = time_split.group(2).strip()
                    else:
                        # 如果精确匹配失败，尝试简单分割
                        time_parts = re.split(r'\s*[–—-]\s*', time_range)
                        if len(time_parts) >= 2:
                            education_item['start_time'] = time_parts[0].strip()
                            education_item['end_time'] = time_parts[1].strip()
                        else:
                            education_item['time_range'] = time_range
                
                # 提取学校
                school_match = re.search(r'学校[：:]\s*([^\n]+)', section)
                if school_match:
                    education_item['school'] = school_match.group(1).strip()
                
                # 提取学历
                degree_match = re.search(r'学历[：:]\s*([^\n]+)', section)
                if degree_match:
                    education_item['degree'] = degree_match.group(1).strip()
                
                # 提取专业
                major_match = re.search(r'专业[：:]\s*([^\n]+)', section)
                if major_match:
                    education_item['major'] = major_match.group(1).strip()
                
                if education_item:
                    education_list.append(education_item)
            
            if education_list:
                return json.dumps(education_list, ensure_ascii=False)
            
        except Exception as e:
            print(f"解析教育经历时出错: {str(e)}")
        
        # 如果解析失败，返回原始文本
        return section_content
    
    def _parse_table_education_experience(self, text: str) -> Optional[str]:
        """解析表格格式的教育经历"""
        try:
            education_list = []
            
            # 查找表格格式的教育经历时间
            # 匹配 "起止时间: 2015-09 – 2019-06" 这样的格式
            time_pattern = r'起止时间[：:]\s*(\d{4}-\d{2})\s*[–—-]\s*(\d{4}-\d{2})'
            time_matches = re.findall(time_pattern, text)
            
            if time_matches:
                for start_time, end_time in time_matches:
                    education_item = {
                        'start_time': start_time,
                        'end_time': end_time
                    }
                    
                    # 在时间附近查找学校信息
                    time_context = f'{start_time}.*?{end_time}'
                    context_match = re.search(f'{time_context}.*?(?=起止时间|工作经历|$)', text, re.DOTALL)
                    if context_match:
                        context = context_match.group(0)
                        
                        # 提取学校
                        school_match = re.search(r'学校[：:]\s*([^\s\n]+)', context)
                        if school_match:
                            education_item['school'] = school_match.group(1).strip()
                        
                        # 提取学历
                        degree_match = re.search(r'学历[：:]\s*([^\s\n]+)', context)
                        if degree_match:
                            education_item['degree'] = degree_match.group(1).strip()
                        
                        # 提取专业
                        major_match = re.search(r'专业[：:]\s*([^\s\n]+)', context)
                        if major_match:
                            education_item['major'] = major_match.group(1).strip()
                    
                    education_list.append(education_item)
                
                if education_list:
                    return json.dumps(education_list, ensure_ascii=False)
            
        except Exception as e:
            print(f"解析表格教育经历时出错: {str(e)}")
        
        return None
    
    def _parse_work_experience(self, text: str, patterns: list) -> Optional[str]:
        """解析工作经历为JSON格式"""
        section_content = self._extract_section_content(text, patterns)
        if not section_content:
            # 尝试从整个文本中提取表格格式的工作经历
            return self._parse_table_work_experience(text)
        
        try:
            work_list = []
            
            # 按段落分割工作经历
            sections = re.split(r'\n\s*\n|\n(?=起止时间)', section_content)
            
            for section in sections:
                if not section.strip():
                    continue
                
                work_item = {}
                
                # 提取起止时间
                time_match = re.search(r'起止时间[：:]\s*([^\n]+)', section)
                if time_match:
                    time_range = time_match.group(1).strip()
                    # 使用更精确的正则来分割时间范围，避免分割日期中的连字符
                    # 匹配类似 "2021-07 – 至今" 的格式
                    time_split = re.match(r'(\d{4}-\d{2})\s*[–—-]\s*(\d{4}-\d{2}|\S+)', time_range)
                    if time_split:
                        work_item['start_time'] = time_split.group(1).strip()
                        work_item['end_time'] = time_split.group(2).strip()
                    else:
                        # 如果精确匹配失败，尝试简单分割
                        time_parts = re.split(r'\s*[–—-]\s*', time_range)
                        if len(time_parts) >= 2:
                            work_item['start_time'] = time_parts[0].strip()
                            work_item['end_time'] = time_parts[1].strip()
                        else:
                            work_item['time_range'] = time_range
                
                # 提取公司
                company_match = re.search(r'公司[：:]\s*([^\n]+)', section)
                if company_match:
                    work_item['company'] = company_match.group(1).strip()
                
                # 提取职位
                position_match = re.search(r'职位[：:]\s*([^\n]+)', section)
                if position_match:
                    work_item['position'] = position_match.group(1).strip()
                
                # 提取主要职责
                responsibility_match = re.search(r'主要职责[：:]\s*([^\n]*(?:\n(?!起止时间|公司|职位)[^\n]*)*)', section, re.MULTILINE)
                if responsibility_match:
                    responsibilities = responsibility_match.group(1).strip()
                    # 清理职责描述
                    responsibilities = re.sub(r'\n+', ' ', responsibilities)
                    work_item['responsibilities'] = responsibilities
                
                if work_item:
                    work_list.append(work_item)
            
            if work_list:
                return json.dumps(work_list, ensure_ascii=False)
            
        except Exception as e:
            print(f"解析工作经历时出错: {str(e)}")
        
        # 如果解析失败，返回原始文本
        return section_content
    
    def _parse_table_work_experience(self, text: str) -> Optional[str]:
        """解析表格格式的工作经历"""
        try:
            work_list = []
            
            # 查找所有工作经历时间，包括"至今"
            time_patterns = [
                r'起止时间[：:]\s*(\d{4}-\d{2})\s*[–—-]\s*(\d{4}-\d{2})',  # 完整时间段
                r'起止时间[：:]\s*(\d{4}-\d{2})\s*[–—-]\s*(至今)'  # 至今
            ]
            
            all_matches = []
            for pattern in time_patterns:
                matches = re.findall(pattern, text)
                all_matches.extend(matches)
            
            if all_matches:
                for start_time, end_time in all_matches:
                    work_item = {
                        'start_time': start_time,
                        'end_time': end_time
                    }
                    
                    # 在时间附近查找公司和职位信息
                    # 查找这个时间段对应的上下文
                    time_context_pattern = f'{re.escape(start_time)}.*?{re.escape(end_time)}'
                    context_match = re.search(f'{time_context_pattern}.*?(?=起止时间|兴趣爱好|自我评价|$)', text, re.DOTALL)
                    
                    if context_match:
                        context = context_match.group(0)
                        
                        # 提取公司
                        company_match = re.search(r'公司[：:]\s*([^\s\n]+)', context)
                        if company_match:
                            work_item['company'] = company_match.group(1).strip()
                        
                        # 提取职位
                        position_match = re.search(r'职位[：:]\s*([^\s\n]+)', context)
                        if position_match:
                            work_item['position'] = position_match.group(1).strip()
                        
                        # 提取主要职责
                        responsibility_match = re.search(r'主要职责[：:]\s*([^起止时间]*?)(?=起止时间|$)', context, re.DOTALL)
                        if responsibility_match:
                            responsibilities = responsibility_match.group(1).strip()
                            # 清理职责描述
                            responsibilities = re.sub(r'\n+', ' ', responsibilities)
                            responsibilities = re.sub(r'\s+', ' ', responsibilities)
                            work_item['responsibilities'] = responsibilities
                    
                    work_list.append(work_item)
                
                if work_list:
                    return json.dumps(work_list, ensure_ascii=False)
            
        except Exception as e:
            print(f"解析表格工作经历时出错: {str(e)}")
        
        return None 